import sys
import openpyxl
import tkinter as tk
from tkinter import ttk, IntVar, StringVar
import tkinter.messagebox as mb
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
from glob import glob

VERSIONS_FOLDER = 'xl_versions'
RESULTS_FOLDER = 'תוצאות'
TEMPLATE_FILE = 'template.docx'
XL_NAME = 'תורנים.xlsx'
TABLE_SIZE = 9 # x 9
SHISHI_INDEX = (1, 4)

# regular, shishi
TABLE_CELLS = '''\
15 16 17 24 25
11 12 20 21
42 43 44 51 52
38 39 47 48
36 37 45 46
67 68 76 77
65 66 74 75
63 64 72 73

13 14 22 23
40 41 49 50'''.split('\n')

# regular, shishi
KAITZ_CELLS = '''\
15 16 24 25
11 12 20
42 43 51 52
38 39 47
36 37 45
67 68 76
65 66 74
63 64 72

13 14 22
40 41 49'''.split('\n')

shiurim = {
    'א': '1',
    'ב': '2',
    'ג': '3',
    'ד': '4',
    'ה': '5',
}

REGULAR = MITUTA = INFO = ZMAN_HOREF = 0
SHISHI = HAVRUTA = ERROR = ZMAN_KAITZ = 1

HOREF_NUM = 4
KAITZ_NUM = 3

class Excel:
    def __init__(self):
        self.wb = openpyxl.load_workbook(XL_NAME)
        self.ws = self.wb.active
        self.toranim_data = {} # {name: [all, shishi]}
        self.havruta_data = []

    def is_empty_cell(self, cell):
        return cell.value is None

    def has_empty_cell(self, range):
        for cell in range:
            if self.is_empty_cell(cell):
                return True

        return False
    
    def extract(self):
        last_mituta = f'D{self.ws.max_row}'
        if self.is_empty_cell(self.ws[last_mituta]):
            Errors.fatal("יש בעיה באקסל, תבדוק שהוא ערוך כמו שצריך. כנראה שהבעיה היא שהכנסת לאקסל שורות נוספות מתחת לתלמיד האחרון בטור של השמות")

        first_mituta = 'A2'
        for person in self.ws[f'{first_mituta}:{last_mituta}']:
            cell_coordinate = person[0].coordinate
            if self.has_empty_cell(person):
                Errors.fatal(f"יש בעיה באקסל, בשורה של התא\n {cell_coordinate}\n יש תא ריק. תמלא אותו ותנסה שוב")

            name = person[0].value
            regular_done = person[2].value
            shishi_done = person[3].value
            self.toranim_data[name] = [regular_done, shishi_done]
        
        first_havruta = 'F2'
        last_havruta = f'G{self.ws.max_row}'
        for havruta in self.ws[f'{first_havruta}:{last_havruta}']:
            if self.is_empty_cell(havruta[0]):
                continue
            self.havruta_data.append([havruta[0].value, havruta[1].value])
            
        return self.havruta_data, self.toranim_data

    def get_havruta(self, name):
        for i in self.havruta_data:
            if name == i[0]:
                return i[1]
            elif name == i[1]:
                return i[0]
        return ''

    def get_sorted_row(self, name, shiur):
        max_row = self.ws.max_row
        for row in range(2, max_row + 1): # first toran row is 2
            row_name = self.ws[f'A{row}'].value
            row_shiur = str(self.ws[f'B{row}'].value)            
            if row_shiur > shiur or (row_shiur == shiur and row_name > name):
                break
        else:
            row = max_row + 1
        self.ws.insert_rows(row)
        return row

    def add_person(self, name, shiur):
        last_person_name = f'A{self.ws.max_row}'
        if self.is_empty_cell(self.ws[last_person_name]):
            Errors.fatal("יש בעיה באקסל, תבדוק שהוא ערוך כמו שצריך. כנראה שהבעיה היא שהכנסת לאקסל שורות נוספות מתחת לתלמיד האחרון בטור של השמות")
        if shiur not in shiurim.keys():
            Errors.warning("אנא הכנס שיעור בין א-ה")
            return
        
        def key_func(cell):
            if isinstance(cell.value, int) or (isinstance(cell.value, str) and cell.value.isdigit()):
                return cell.value
            else:
                Errors.fatal(f"יש טעות באקסל במספר התורנויות בתא\n{cell.coordinate}\nתקן את הטעות ונסה שנית")
        max_regular_cell = max(self.ws['C'][1:], key=key_func)
        max_shishi_cell = max(self.ws['D'][1:], key=key_func)
        
        shiur_number = shiurim[shiur]
        row_to_insert = self.get_sorted_row(name, shiur_number)

        self.ws[f'A{row_to_insert}'] = name
        self.ws[f'B{row_to_insert}'] = shiur_number
        self.ws[f'C{row_to_insert}'] = max_regular_cell.value
        self.ws[f'D{row_to_insert}'] = max_shishi_cell.value

        self.wb.save(XL_NAME)
        Tkinter.show(INFO, f'הוספת את {name} משיעור {shiur} בהצלחה')
    
    def update(self, toranim_data):
        if not os.path.exists(VERSIONS_FOLDER):
            os.mkdir(VERSIONS_FOLDER)
        self.wb.save(f'{VERSIONS_FOLDER}/{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.xlsx')
        for i, name in enumerate(toranim_data):
            self.ws.cell(i+2, 3).value = toranim_data[name][REGULAR]
            self.ws.cell(i+2, 4).value = toranim_data[name][SHISHI]
        self.wb.save(XL_NAME)

class Tkinter:        
    zman = None
    strvar_nums = []

    @classmethod
    def is_special_sevev(cls):
        return cls.strvar_nums

    @classmethod
    def is_zman_horef(cls):
        return cls.zman.get() == ZMAN_HOREF

    @classmethod
    def start(cls):
        cls.root = tk.Tk()
        cls.root.title('שיבוץ תורנים')
        frame = tk.Frame(cls.root, padx=60, pady=60)
        frame.pack()

        cls.zman = IntVar(value=ZMAN_HOREF)
        ttk.Radiobutton(frame, text="זמן אלול / חורף", value=ZMAN_HOREF, var=cls.zman).pack(anchor='w')
        ttk.Radiobutton(frame, text="זמן קיץ", value=ZMAN_KAITZ, var=cls.zman).pack(anchor='w')

        ttk.Button(frame, text='סבב רגיל', command=lambda: Calculate().calculate()).pack()
        ttk.Button(frame, text='סבב מיוחד', command=cls.special_sevev).pack()
        ttk.Button(frame, text='שחזר פעם אחרונה', command=cls.restore).pack()
        ttk.Button(frame, text='הוסף בחור ישיבה', command=cls.add_person_dialog).pack()
        cls.root.mainloop()

    @classmethod
    def add_person_dialog(cls):
        dialog = tk.Toplevel(cls.root)
        dialog.title("הוסף בחור ישיבה")
        dialog.geometry("300x150")

        name_var = StringVar()
        shiur_var = StringVar()

        name_label = ttk.Label(dialog, text="הכנס שם מלא:")
        name_label.grid(column=0, row=0, padx=5, pady=5)
        name_entry = ttk.Entry(dialog, textvariable=name_var)
        name_entry.grid(column=1, row=0, padx=5, pady=5)

        options = list(shiurim.keys())
        shiur_var.set(options[0])
        shiur_label = ttk.Label(dialog, text="בחר שיעור עבור הבחור:")
        shiur_label.grid(column=0, row=1, padx=5, pady=5)
        shiur_dropdown = ttk.OptionMenu(dialog, shiur_var, *options)
        shiur_dropdown.grid(column=1, row=1, padx=5, pady=5)

        def submit_dialog():
            name = name_var.get()
            shiur = shiur_var.get()
            if len(name.split()) < 2:
                Errors.warning("נא להכניס שתי מילים לפחות בשם המלא כדי שהתורן יוכל לדעת שזה הוא")
            elif shiur not in shiurim.keys():
                Errors.warning("נא לא להכניס בחור ישיבה משיעור שאינו מהשיעורים א-ה") 
            else:
                dialog.destroy()
                Excel().add_person(name, shiur)

        submit_button = ttk.Button(dialog, text="תכניס אותו לתורנות", command=submit_dialog)
        submit_button.grid(column=0, row=2, columnspan=2, padx=5, pady=5)

        dialog.transient(cls.root)
        dialog.grab_set()
        cls.root.wait_window(dialog)

    @classmethod
    def remove_frame(cls):
        cls.root.winfo_children()[0].destroy()

    @classmethod
    def get_int_counts(cls):
        return [int(i.get()) for i in cls.strvar_nums if i]
    
    @classmethod
    def special_sevev(cls):
        cls.remove_frame()
        frame = tk.Frame(cls.root, padx=60, pady=60)
        frame.pack()

        DAYS = ['ראשון-שני', 'שלישי-רביעי', 'חמישי-שישי', 'שבת']
        cls.strvar_nums = [0] * 12
        for i in range(4):
            tk.Label(frame,text=DAYS[3-i]).grid(row=0, column=i)
            for j in range(3):
                if i == 0 and j == 2 or i == 3 and j == 0:
                    continue
                frame1 = tk.Frame(frame)
                frame1.grid(row=j+1, column=i)
                cls.strvar_nums[j*4+i] = tk.StringVar()
                toranim_num = HOREF_NUM if Tkinter.is_zman_horef() else KAITZ_NUM
                cls.strvar_nums[j*4+i].set(str(toranim_num + int(i==0)))
                spinbox=tk.Spinbox(frame1, from_=0, to=50, textvariable=cls.strvar_nums[j*4+i])
                spinbox.pack()
    
        ttk.Button(cls.root, text='סבבה', command=lambda:
         Calculate().calculate()).pack()
    
    @classmethod
    def get_sums(cls):
        if cls.is_special_sevev():
            int_nums = cls.get_int_counts()
            shishis_sum = sum([int_nums[i] for i in SHISHI_INDEX])
            rest_sum = sum(int_nums) - shishis_sum
            return [rest_sum, shishis_sum]
        elif cls.zman.get() == ZMAN_KAITZ:
            return [26, 6] # rest: ((3 + 3 + 4) * 3 - 4), shishis: (3 * 2)
        else:
            return [34, 8] # rest: ((4 + 4 + 5) * 3 - 5), shishis: (4 * 2)

    @classmethod
    def close(cls):
        cls.root.destroy()

    @classmethod
    def show(cls, type, msg):
        if type == ERROR:
            mb.showerror('קרתה תקלה', msg)
        else:
            mb.showinfo('כל הכבוד', msg)

    @classmethod
    def restore(cls):
        xl_folder, word_foler = glob(VERSIONS_FOLDER + '/*'), glob(RESULTS_FOLDER + '/*')
        if not xl_folder or not word_foler:
            Errors.warning("אין לך מה לשחזר")
        else:
            os.remove(XL_NAME)
            os.remove(max(word_foler, key=os.path.getctime)) # latest file in folder
            os.rename(max(xl_folder, key=os.path.getctime), XL_NAME)
            cls.show(INFO, 'שוחזר בהצלחה')

class Calculate:
    def __init__(self):
        self.excel = Excel()
        self.havrutot, self.toranim = self.excel.extract()
        self.results = [[[], []], [[], []]] # [regular: [lonely, havruta], shishi: [lonely, havruta]]
        self.count = [0, 0] # regular and shishi
        self.min_lists = [[], []] # regular and shishi

    def get_min_list(self, type):
        res = []
        min_shishi = min([i[SHISHI] for i in self.toranim.values()])
        min_regular = min([i[REGULAR] for i in self.toranim.values()])
        for i in self.toranim:
            if type == SHISHI:
                if self.toranim[i][REGULAR] == min_regular and self.toranim[i][SHISHI] == min_shishi:
                    res.append(i)
            else:
                if self.toranim[i][REGULAR] == min_regular:
                    res.append(i)
        if not res: # if there's no one in united min list
            for i in self.toranim:
                if self.toranim[i][SHISHI] == min_shishi:
                    res.append(i)
        return res

    def add(self, name, has_havruta, type):
        self.results[type][has_havruta].append(name)
        self.count[type] -= 1
        self.toranim[name][REGULAR] += 1
        if type == SHISHI:
            self.toranim[name][SHISHI] += 1
        if len(self.min_lists[type]) == 0:
            Errors.warning(f"ייתכן שיש בעיה עם {name}, בדוק ליתר ביטחון")
        else:
            self.min_lists[type].remove(name)

    def add_last_toran(self, type):
        for name in self.get_min_list(type):
            if self.excel.get_havruta(name) == '':
                self.add(name, MITUTA, type)
                return True
        return False

    def util1(self, type):
        name = self.min_lists[type][0]
        if self.count[type] == 1:
            self.add_last_toran(type)
            return
        havruta = self.excel.get_havruta(name)
        if havruta in self.min_lists[type]:
            self.add(havruta, HAVRUTA, type)
            self.add(name, HAVRUTA, type)
        else:
            self.add(name, MITUTA, type)

    def get_odd_count(self):
        if Tkinter.is_special_sevev():
            nums = Tkinter.get_int_counts()
            self.odds_count = sum([i % 2 for i in nums])
        elif Tkinter.is_zman_horef():
            self.odds_count = 2 # shabat = 5, 2 shabattot
        else:
            self.odds_count = 8 # everyday = 8, 9 times

    def util(self, type):
        self.min_lists[type] = self.get_min_list(type)
        if len(self.min_lists[type]) < self.count[type]:
            for _ in self.min_lists[type]:
                self.util1(type)
            self.min_lists[type] = self.get_min_list(type)
        while self.count[type] > 0:
            self.util1(type)
            # odds_count = days that need a single toran
            if self.count[type] and len(self.results[type][MITUTA]) < self.odds_count:
                self.add_last_toran(type)
            self.min_lists[type] = self.get_min_list(type)

    def calculate(self):
        self.count = Tkinter.get_sums()
        self.get_odd_count()
        self.util(SHISHI)
        self.util(REGULAR)
        self.save_results()
        Tkinter.show(INFO,
        'התוצאה שמורה ב"תוצאות" והקובץ "תורנים" התעדכן. אם אתה רוצה להתחרט אז תיכנס שוב לתוכנה ותלחץ על "שחזר".')
        Tkinter.close()

    def save_results(self):
        word = Word()
        word.update_table_cells()
        word.fill_table(self.results)
        word.save()
        self.excel.update(self.toranim)

class Word:
    def __init__(self):
        self.doc = Document('template.docx')
        self.table = self.doc.tables[0]

    def update_table_cells(self):
        self.table_cells = [[], []] # regular, shishi
        index = REGULAR

        index_table_cells = TABLE_CELLS if Tkinter.is_zman_horef() else KAITZ_CELLS
        for i in index_table_cells:
            if i == '':
                index = SHISHI
                continue
            self.table_cells[index].append([int(i) for i in i.split(' ')])

        if Tkinter.strvar_nums: # special sevev
            int_counts = Tkinter.get_int_counts()
            shishi_counts = [int_counts[i] for i in SHISHI_INDEX]
            rest_counts = []
            for key, value in enumerate(int_counts):
                if key not in SHISHI_INDEX:
                    rest_counts.append(value)
            counts = [rest_counts, shishi_counts]
            for i in range(len(self.table_cells)): # regular and shishi
                for j in range(len(self.table_cells[i])): # groups
                    if counts[i][j] < len(self.table_cells[i][j]):
                        self.table_cells[i][j] = self.table_cells[i][j][:counts[i][j]]
                    else:
                        self.table_cells[i][j] = (self.table_cells[i][j] * (counts[i][j] //
                         len(self.table_cells[i][j]) + 1))[:counts[i][j]]

    def fill_table(self, results):
        counts = [[0, 0], [0, 0]] # regular [lonely, havruta], shishi [lonely, havruta]
        for i in range(2): # 2 iterations - regular and shishi
            for j in self.table_cells[i]: # over groups.
                flag = False
                for iteration, k in enumerate(j):
                    cell = self.table.rows[k // TABLE_SIZE].cells[k % TABLE_SIZE]
                    if cell.text:
                        cell.text += ',\n'
                    # flag is True when one of havruta was inserted
                    if ( counts[i][HAVRUTA] < len(results[i][HAVRUTA]) and iteration < len(j) - 1 ) or flag:
                        cell.text += results[i][HAVRUTA][counts[i][HAVRUTA]]
                        flag = not flag
                        counts[i][HAVRUTA] += 1
                    else:
                        cell.text += results[i][MITUTA][counts[i][MITUTA]]
                        counts[i][MITUTA] += 1
                    cell.paragraphs[0].runs[0].font.size = Pt(12)

    def save(self):
        if not os.path.exists(RESULTS_FOLDER):
            os.mkdir(RESULTS_FOLDER)
        self.doc.save(f'{RESULTS_FOLDER}/{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx')

class Errors:
    @classmethod
    def fatal(cls, msg):
        Tkinter.show(ERROR, msg)
        Tkinter.close()
        sys.exit(0)

    @classmethod
    def warning(cls, msg):
        Tkinter.show(ERROR, msg)
        

if __name__ == '__main__':
    if not os.path.exists(TEMPLATE_FILE) or not os.path.exists(XL_NAME):
        Errors.fatal('חסרים לך קבצים (צריך שיהיה לך: template.docx, תורנים.xlsx)')
    else:
        Tkinter.start()
