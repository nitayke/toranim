import openpyxl
import tkinter as tk
from tkinter import ttk
from docx import Document
from docx.shared import Pt
from datetime import datetime

FOLDER = 'C:/Users/user/Desktop/toranim/'
VERSIONS_FOLDER = 'xl_versions'
RESULTS_FOLDER = 'תוצאות'
XL_NAME = 'try.xlsx'
TABLE_SIZE = 9 # x 9
# regular, shishi
TABLE_CELLS = '''15, 16, 17, 24, 25
11, 12, 20, 21
9, 10, 18, 19
42, 43, 44, 51, 52
38, 39, 47, 48
36, 37, 45, 46
67, 68, 76, 77
65, 66, 74, 75
63, 64, 72, 73

13, 14, 22, 23
40, 41, 49, 50'''.split('\n')

REGULAR = MITUTA = 0
SHISHI = HAVRUTA = 1


class Excel:
    def __init__(self):
        self.wb = openpyxl.load_workbook(FOLDER + XL_NAME)
        self.ws = self.wb.active
        self.toranim_data = {} # {name: [all, shishi]}
        self.havruta_data = []
    
    def extract(self):
        for i in self.ws['A2:D' + str(self.ws.max_row)]:
            self.toranim_data[i[0].value] = [i[2].value, i[3].value] # regular and shishi
        
        for i in self.ws['F2:G' + str(self.ws.max_row // 2)]:
            if not i[0].value:
                break
            self.havruta_data.append([i[0].value, i[1].value])
        return self.havruta_data, self.toranim_data

    def get_havruta(self, name):
        for i in self.havruta_data:
            if name == i[0]:
                return i[1]
            elif name == i[1]:
                return i[0]
        return ''
    
    def update(self, toranim_data):
        # self.wb.save(f'{FOLDER}{VERSIONS_FOLDER}/{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.xlsx')
        for i, name in enumerate(toranim_data):
            self.ws.cell(i+2, 3).value = toranim_data[name][REGULAR]
            self.ws.cell(i+2, 4).value = toranim_data[name][SHISHI]
        self.wb.save(FOLDER + XL_NAME)


class Tkinter:
    root = tk.Tk()
    strvar_nums = []

    @classmethod
    def start(cls):
        cls.root.title('שיבוץ תורנים')
        frame = tk.Frame(cls.root, padx=60, pady=60)
        frame.pack()
        ttk.Button(frame, text='סבב רגיל', command=lambda: Calculate().calculate()).pack()
        ttk.Button(frame, text='סבב מיוחד', command=cls.special_sevev).pack()
        cls.root.mainloop()

    @classmethod
    def remove_frame(cls):
        cls.root.winfo_children()[0].destroy()

    @classmethod
    def get_int_counts(cls):
        return [int(i.get()) for i in cls.strvar_nums if i]
    
    @classmethod
    def special_sevev(cls):
        cls.remove_frame()
        frame = tk.Frame(cls.root, padx=60, pady=60)
        frame.pack()

        DAYS = ['ראשון-שני', 'שלישי-רביעי', 'חמישי-שישי', 'שבת']
        cls.strvar_nums = [0] * 12
        for i in range(4):
            tk.Label(frame,text=DAYS[3-i]).grid(row=0, column=i)
            for j in range(3):
                if i == 0 and j == 2:
                    break
                frame1 = tk.Frame(frame)
                frame1.grid(row=j+1, column=i)
                cls.strvar_nums[j*4+i] = tk.StringVar()
                cls.strvar_nums[j*4+i].set(str(4 + int(i==0)))
                spinbox=tk.Spinbox(frame1, from_=0, to=10, textvariable=cls.strvar_nums[j*4+i])
                spinbox.pack()
    
        ttk.Button(cls.root, text='סבבה', command=lambda:
         Calculate().calculate()).pack()
    
    @classmethod
    def get_sums(cls):
        if cls.strvar_nums:
            int_nums = cls.get_int_counts()
            shishis_sum = int_nums[1] + int_nums[5]
            rest_sum = sum(int_nums) - shishis_sum
            return [rest_sum, shishis_sum]
        else:
            return [38, 8]

    @classmethod
    def close(cls):
        cls.root.destroy()


class Calculate:
    def __init__(self):
        self.excel = Excel()
        self.havrutot, self.toranim = self.excel.extract()
        self.results = [[[], []], [[], []]] # [regular: [lonely, havruta], shishi: [lonely, havruta]]
        self.count = [0, 0] # regular and shishi

    def get_min_list(self, shishi):
        res = []
        min_shishi = min([i[SHISHI] for i in self.toranim.values()])
        min_toranut = min([i[REGULAR] for i in self.toranim.values()])
        for i in self.toranim:
            if shishi:
                if self.toranim[i][REGULAR] == min_toranut and self.toranim[i][SHISHI] == min_shishi:
                    res.append(i)
            else:
                if self.toranim[i][REGULAR] == min_toranut:
                    res.append(i)
        return res

    def add(self, name, has_havruta, shishi):
        self.results[shishi][has_havruta].append(name)
        self.count[shishi] -= 1
        self.toranim[name][REGULAR] += 1
        if shishi:
            self.toranim[name][SHISHI] += 1
    
    def add_last_toran(self, shishi):
        for i in self.get_min_list(shishi):
            if self.excel.get_havruta(i) == '':
                self.add(i, False, shishi)
                return True
        return False

    def util(self, shishi):
        if len(self.min_lists[shishi]) <= self.count[shishi]:
            self.results[shishi][MITUTA] = self.min_lists[shishi] # TODO: MAYBE NOT SUPPOSED TO BE LONELY
            self.count[shishi] -= len(self.min_lists[shishi])  # need to add more to fill the required
            for i in self.min_lists[shishi]:
                self.toranim[i][shishi] += 1
        for i in self.min_lists[shishi]:
            if self.count[shishi] == 1 and not self.add_last_toran(shishi):
                self.add(i, False, shishi)
            if self.count[shishi] == 0:
                break
            current_havruta = self.excel.get_havruta(i)
            if current_havruta in self.min_lists[shishi]:
                self.add(current_havruta, True, shishi)
                self.add(i, True, shishi)
            else:
                self.add(i, False, shishi)

    def calculate(self): # TODO: FIX IT
        self.count = Tkinter.get_sums()
        self.min_lists = [[], self.get_min_list(SHISHI)] # regular and united
        # united is where the regular and shishi have minimum values
        self.util(SHISHI)
        self.min_lists[REGULAR] = self.get_min_list(REGULAR)
        self.util(REGULAR)
        self.save_results()
        self.excel.update(self.toranim)
        Tkinter.close()

    def save_results(self):
        print('results\n', self.results)
        word = Word()
        word.update_table_cells()
        word.fill_table(self.results)
        word.save()


class Word:
    def __init__(self):
        self.doc = Document(FOLDER + 'template.docx')
        self.table = self.doc.tables[0]

    def update_table_cells(self):
        self.table_cells = [[], []] # regular, shishi
        index = 0
        for i in TABLE_CELLS:
            if i == '':
                index = 1
                continue
            self.table_cells[index].append([int(i) for i in i.split(', ')])
        if Tkinter.strvar_nums: # special sevev
            int_counts = Tkinter.get_int_counts()
            shishi_counts = [int_counts[1], int_counts[5]]
            tmp = shishi_counts.copy()
            rest_counts = [i for i in int_counts if not i in tmp or tmp.remove(i)]
            counts = [rest_counts, shishi_counts]
            for i in range(len(self.table_cells)): # regular and shishi
                for j in range(len(self.table_cells[i])): # groups
                    if counts[i][j] < len(self.table_cells[i][j]):
                        self.table_cells[i][j] = self.table_cells[i][j][:counts[i][j]]
                    else:
                        self.table_cells[i][j] = (self.table_cells[i][j] * (counts[i][j] //
                         len(self.table_cells[i][j]) + 1))[:counts[i][j]]
            

    
    def fill_table(self, results):
        counts = [[0, 0], [0, 0]] # regular [lonely, havruta], shishi [lonely, havruta]

        lengths = ((len(results[REGULAR][MITUTA]), len(results[REGULAR][HAVRUTA])),
            ((len(results[SHISHI][MITUTA]), len(results[SHISHI][HAVRUTA]))))
        
        for i in range(len(self.table_cells)): # 2 iterations - regular and shishi
            for j in self.table_cells[i]: # over groups.
                flag = False
                counter = 0
                for k in j:
                    cell = self.table.rows[k // TABLE_SIZE].cells[k % TABLE_SIZE]
                    if cell.text:
                        cell.text += ',\n'
                    if counts[i][HAVRUTA] < lengths[i][HAVRUTA] and counter < len(j) - 1 or flag:
                        cell.text += results[i][HAVRUTA][counts[i][HAVRUTA]]
                        flag = not flag
                        counts[i][HAVRUTA] += 1
                    else:
                        cell.text += results[i][MITUTA][counts[i][MITUTA]]
                        counts[i][MITUTA] += 1
                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                    counter += 1

    def save(self):
        self.doc.save(f'{FOLDER}{RESULTS_FOLDER}/{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx')



if __name__ == '__main__':
    Tkinter.start()