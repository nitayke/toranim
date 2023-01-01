import openpyxl
import tkinter as tk
from tkinter import ttk
import tkinter.messagebox as mb
from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
from glob import glob

VERSIONS_FOLDER = 'xl_versions'
RESULTS_FOLDER = 'תוצאות'
TEMPLATE_FILE = 'template.docx'
XL_NAME = 'תורנים.xlsx'
TABLE_CELLS = '''2 3
4 5
8 9
10 11'''.split('\n')

MITUTA = INFO = 0
HAVRUTA = ERROR = 1


class Excel:
    def __init__(self):
        self.wb = openpyxl.load_workbook(XL_NAME)
        self.ws = self.wb.active
        self.toranim_data = {} # {name: count}
        self.havruta_data = []
    
    def extract(self):
        for i in self.ws['A2:C' + str(self.ws.max_row)]:
            self.toranim_data[i[0].value] = i[2].value
        
        for i in self.ws['F2:G' + str(self.ws.max_row // 2)]:
            if not i[0].value:
                break
            self.havruta_data.append([i[0].value, i[1].value])
        return self.havruta_data, self.toranim_data

    def get_havruta(self, name):
        for i in self.havruta_data:
            if name == i[0]:
                return i[1]
            elif name == i[1]:
                return i[0]
        return ''
    
    def update(self, toranim_data):
        if not os.path.exists(VERSIONS_FOLDER):
            os.mkdir(VERSIONS_FOLDER)
        self.wb.save(f'{VERSIONS_FOLDER}/{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.xlsx')
        for i, name in enumerate(toranim_data):
            self.ws.cell(i+2, 3).value = toranim_data[name]
        self.wb.save(XL_NAME)


class Tkinter:
    strvar_nums = []

    @classmethod
    def start(cls):
        cls.root = tk.Tk()
        cls.root.title('שיבוץ תורנים')
        frame = tk.Frame(cls.root, padx=60, pady=60)
        frame.pack()
        ttk.Button(frame, text='שבץ', command=lambda: Calculate().calculate()).pack()
        ttk.Button(frame, text='שחזר פעם אחרונה', command=cls.restore).pack()
        cls.root.mainloop()

    @classmethod
    def remove_frame(cls):
        cls.root.winfo_children()[0].destroy()

    @classmethod
    def get_int_counts(cls):
        return [int(i.get()) for i in cls.strvar_nums if i]
    
    @classmethod
    def close(cls):
        cls.root.destroy()

    @classmethod
    def show(cls, type, msg):
        if type == ERROR:
            mb.showerror('קרתה תקלה', msg)
        else:
            mb.showinfo('כל הכבוד', msg)

    @classmethod
    def restore(cls):
        xl_folder, word_foler = glob(VERSIONS_FOLDER + '/*'), glob(RESULTS_FOLDER + '/*')
        if not xl_folder or not word_foler:
            cls.show(ERROR, "אין לך מה לשחזר")
        else:
            os.remove(XL_NAME)
            os.remove(max(word_foler, key=os.path.getctime)) # latest file in folder
            os.rename(max(xl_folder, key=os.path.getctime), XL_NAME)
            cls.show(INFO, 'שוחזר בהצלחה')


class Calculate:
    def __init__(self):
        self.excel = Excel()
        self.havrutot, self.toranim = self.excel.extract()
        self.results = [[], []] # [lonely, havruta]
        self.count = 0
        self.min_list = []

    def get_min_list(self):
        res = []
        minimum = min(self.toranim.values())
        for i in self.toranim:
            if self.toranim[i] == minimum:
                res.append(i)
        return res

    def add(self, name, has_havruta):
        self.results[has_havruta].append(name)
        self.count -= 1
        self.toranim[name] += 1
        self.min_list.remove(name)

    def add_last_toran(self):
        for i in self.get_min_list():
            if self.excel.get_havruta(i) == '':
                self.add(i, MITUTA)
                return True
        return False

    def util1(self):
        name = self.min_list[0]
        if self.count == 1:
            self.add_last_toran()
            return
        havruta = self.excel.get_havruta(name)
        if havruta in self.min_list:
            self.add(havruta, HAVRUTA)
            self.add(name, HAVRUTA)
        else:
            self.add(name, MITUTA)

    def get_odd_count(self):
        if Tkinter.strvar_nums:
            nums = Tkinter.get_int_counts()
            self.odds_count = sum([i % 2 for i in nums])
        else:
            self.odds_count = 2

    def util(self):
        self.min_list = self.get_min_list()
        if len(self.min_list) < self.count:
            for i in self.min_list:
                self.util1()
            self.min_list = self.get_min_list()
        while self.count > 0:
            self.util1()
            # odds_count = days that need a single toran
            if self.count and len(self.results[MITUTA]) < self.odds_count:
                self.add_last_toran()
            self.min_list = self.get_min_list()

    def calculate(self):
        self.count = 8
        self.get_odd_count()
        self.util()
        self.util()
        self.save_results()
        Tkinter.show(INFO,
        'התוצאה שמורה ב"תוצאות" והקובץ "תורנים" התעדכן. אם אתה רוצה להתחרט אז תיכנס שוב לתוכנה ותלחץ על "שחזר".')
        Tkinter.close()

    def save_results(self):
        word = Word()
        word.update_table_cells()
        word.fill_table(self.results)
        word.save()
        self.excel.update(self.toranim)


class Word:
    def __init__(self):
        self.doc = Document('template.docx')
        self.table = self.doc.tables[0]

    def update_table_cells(self):
        self.table_cells = []
        for i in TABLE_CELLS:
            self.table_cells.append([int(i) for i in i.split(' ')])

    def fill_table(self, results):
        counts = [0, 0] # [lonely, havruta]
        print(self.table_cells)
        for j in self.table_cells: # over groups.
            flag = False
            for iteration, k in enumerate(j):
                cell = self.table.rows[k // 2].cells[k % 2]
                if cell.text:
                    cell.text += ',\n'
                # flag is True when one of havruta was inserted
                if ( counts[HAVRUTA] < len(results[HAVRUTA]) and iteration < len(j) - 1 ) or flag:
                    cell.text += results[HAVRUTA][counts[HAVRUTA]]
                    flag = not flag
                    counts[HAVRUTA] += 1
                else:
                    cell.text += results[MITUTA][counts[MITUTA]]
                    counts[MITUTA] += 1
                cell.paragraphs[0].runs[0].font.size = Pt(12)

    def save(self):
        if not os.path.exists(RESULTS_FOLDER):
            os.mkdir(RESULTS_FOLDER)
        self.doc.save(f'{RESULTS_FOLDER}/{datetime.now().strftime("%d-%m-%Y_%H-%M-%S")}.docx')


if __name__ == '__main__':
    if not os.path.exists(TEMPLATE_FILE) or not os.path.exists(XL_NAME):
        Tkinter.show(ERROR, 'חסרים לך קבצים (צריך שיהיה לך: template.docx, תורנים.xlsx)')
    else:
        Tkinter.start()
